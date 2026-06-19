#!/usr/bin/env python3
"""文档格式检测脚本 for Paper Formatter Skill.

分析 .docx 文件的格式参数，生成可读的格式摘要报告。
输出 JSON 格式便于后续程序处理。
"""

import json
import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Pt, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误：需要安装 python-docx。运行：python scripts/install_deps.py")
    sys.exit(1)


# 对齐方式映射
ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
    WD_ALIGN_PARAGRAPH.CENTER: "居中",
    WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "分散对齐",
}

# 纸张大小映射（近似）
PAPER_SIZES = {
    (21.0, 29.7): "A4",
    (21.6, 27.9): "Letter",
    (21.6, 35.6): "Legal",
    (18.4, 26.0): "B5",
}


def emu_to_cm(emu) -> float:
    """EMU 转厘米."""
    if emu is None:
        return None
    return round(emu / 360000, 2)


def emu_to_pt(emu) -> float | None:
    """EMU 转磅（pt）."""
    if emu is None:
        return None
    return round(emu / 12700, 1)


def detect_paper_size(width_cm: float, height_cm: float) -> str:
    """根据尺寸推断纸张类型."""
    for (w, h), name in PAPER_SIZES.items():
        if abs(width_cm - w) < 0.5 and abs(height_cm - h) < 0.5:
            return name
    return f"{width_cm:.1f}cm × {height_cm:.1f}cm"


def detect_document_format(filepath: str) -> dict:
    """检测 .docx 文件的格式参数，返回结构化报告."""
    doc = Document(filepath)

    report = {
        "file": os.path.basename(filepath),
        "file_path": os.path.abspath(filepath),
        "page": {},
        "fonts_used": set(),
        "font_sizes_used": set(),
        "heading_styles": {},
        "paragraph_stats": {
            "total": 0,
            "alignment_distribution": {},
            "avg_line_spacing": None,
        },
        "styles_in_use": [],
    }

    # --- 页面设置 ---
    if doc.sections:
        section = doc.sections[0]
        width_cm = emu_to_cm(section.page_width)
        height_cm = emu_to_cm(section.page_height)
        report["page"] = {
            "paper_size": detect_paper_size(width_cm, height_cm),
            "width_cm": width_cm,
            "height_cm": height_cm,
            "margin_top_cm": emu_to_cm(section.top_margin),
            "margin_bottom_cm": emu_to_cm(section.bottom_margin),
            "margin_left_cm": emu_to_cm(section.left_margin),
            "margin_right_cm": emu_to_cm(section.right_margin),
            "orientation": "横向" if width_cm > height_cm else "纵向",
        }

    # --- 段落与字体统计 ---
    line_spacings = []
    alignment_counts = {"左对齐": 0, "居中": 0, "右对齐": 0, "两端对齐": 0, "分散对齐": 0}

    for para in doc.paragraphs:
        report["paragraph_stats"]["total"] += 1

        # 对齐方式
        align_str = ALIGNMENT_MAP.get(para.alignment, "未知")
        if align_str in alignment_counts:
            alignment_counts[align_str] += 1

        # 段落间距
        pf = para.paragraph_format
        ls = pf.line_spacing
        if ls is not None and ls != 1.0:  # 忽略默认值
            line_spacings.append(ls)

        # 字体收集
        for run in para.runs:
            if run.font.name:
                report["fonts_used"].add(run.font.name)
            if run.font.size:
                size_pt = round(run.font.size / 12700, 1)
                report["font_sizes_used"].add(size_pt)

        # 标题样式检测
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name or "heading" in style_name or "标题" in style_name:
            if style_name not in report["heading_styles"]:
                sample_font = "未知"
                sample_size = "未知"
                if para.runs:
                    sample_font = para.runs[0].font.name or "未设置"
                    sample_size = (
                        round(para.runs[0].font.size / 12700, 1)
                        if para.runs[0].font.size
                        else "未设置"
                    )
                report["heading_styles"][style_name] = {
                    "count": 0,
                    "sample_text": para.text[:50] if para.text else "(空)",
                    "font": sample_font,
                    "font_size_pt": sample_size,
                    "bold": para.runs[0].bold if para.runs else None,
                }
            report["heading_styles"][style_name]["count"] += 1

    # 对齐分布（仅保留非零项）
    report["paragraph_stats"]["alignment_distribution"] = {
        k: v for k, v in alignment_counts.items() if v > 0
    }

    # 行距统计
    if line_spacings:
        report["paragraph_stats"]["avg_line_spacing"] = round(
            sum(line_spacings) / len(line_spacings), 2
        )
        report["paragraph_stats"]["line_spacing_range"] = [
            round(min(line_spacings), 2),
            round(max(line_spacings), 2),
        ]

    # --- 已使用的样式 ---
    for style in doc.styles:
        if style.type is not None and style.name:
            report["styles_in_use"].append(style.name)

    # --- 转换 set 为 list ---
    report["fonts_used"] = sorted(list(report["fonts_used"]))
    report["font_sizes_used"] = sorted(list(report["font_sizes_used"]))

    # --- 页眉页脚 ---
    for i, section in enumerate(doc.sections):
        header = section.header
        footer = section.footer
        if header and header.paragraphs:
            header_text = "".join(p.text for p in header.paragraphs if p.text)
            if header_text:
                report.setdefault("headers_footers", {})
                report["headers_footers"][f"section_{i+1}_header"] = header_text[:100]
        if footer and footer.paragraphs:
            footer_text = "".join(p.text for p in footer.paragraphs if p.text)
            if footer_text:
                report.setdefault("headers_footers", {})
                report["headers_footers"][f"section_{i+1}_footer"] = footer_text[:100]

    return report


def print_report(report: dict, verbose: bool = False):
    """以可读格式打印检测报告."""
    print("\n" + "=" * 60)
    print(f"  📄 文档格式检测报告：{report['file']}")
    print("=" * 60)

    # 页面信息
    page = report.get("page", {})
    if page:
        print(f"\n📐 页面设置：")
        print(f"   纸张大小：{page.get('paper_size', '未知')}")
        print(f"   方向：{page.get('orientation', '未知')}")
        print(f"   页边距：上 {page.get('margin_top_cm')}cm 下 {page.get('margin_bottom_cm')}cm "
              f"左 {page.get('margin_left_cm')}cm 右 {page.get('margin_right_cm')}cm")

    # 字体信息
    fonts = report.get("fonts_used", [])
    sizes = report.get("font_sizes_used", [])
    if fonts or sizes:
        print(f"\n🔤 字体概况：")
        print(f"   使用字体：{', '.join(fonts) if fonts else '未检测到'}")
        print(f"   使用字号：{', '.join(str(s) + 'pt' for s in sizes) if sizes else '未检测到'}")

    # 段落统计
    ps = report.get("paragraph_stats", {})
    print(f"\n📝 段落统计：")
    print(f"   总段落数：{ps.get('total', 0)}")
    print(f"   对齐分布：{ps.get('alignment_distribution', {})}")
    if ps.get("avg_line_spacing"):
        print(f"   平均行距：{ps.get('avg_line_spacing')} 倍")

    # 标题样式
    headings = report.get("heading_styles", {})
    if headings:
        print(f"\n📑 标题样式：")
        for name, info in headings.items():
            print(f"   {name}（{info['count']}个）：字体={info['font']} "
                  f"字号={info['font_size_pt']}pt 加粗={info['bold']} ")
            print(f"         示例：{info['sample_text']}")

    # 页眉页脚
    hf = report.get("headers_footers", {})
    if hf:
        print(f"\n📋 页眉页脚：")
        for key, text in hf.items():
            print(f"   {key}: {text}")

    if verbose:
        styles = report.get("styles_in_use", [])
        if styles:
            print(f"\n🎨 已使用样式（前20个）：")
            for s in styles[:20]:
                print(f"   - {s}")

    print("\n" + "=" * 60)


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="检测 .docx 文档格式参数"
    )
    parser.add_argument("file", help="要检测的 .docx 文件路径")
    parser.add_argument(
        "--json", action="store_true", help="以 JSON 格式输出"
    )
    parser.add_argument(
        "--verbose", "-v", action="store_true", help="显示更详细的信息"
    )
    args = parser.parse_args()

    if not os.path.exists(args.file):
        print(f"错误：文件不存在 - {args.file}")
        sys.exit(1)

    report = detect_document_format(args.file)

    if args.json:
        # 转换 set 为 list 用于 JSON 序列化（已在 detect_document_format 中处理）
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print_report(report, verbose=args.verbose)

    # 同时输出 JSON 到 stdout 的后半部分（供程序解析，用分隔符分隔）
    print("\n--- JSON OUTPUT ---")
    print(json.dumps(report, ensure_ascii=False))


if __name__ == "__main__":
    main()
