#!/usr/bin/env python3
"""论文格式修改核心引擎 for Paper Formatter Skill.

支持两种模式：
  1. 指令模式：通过 JSON 配置文件指定格式要求
  2. 模版模式：从模版 .docx 提取样式并应用到目标文档

修改前自动创建 .bak 备份文件。
"""
import io
import sys

# 修复 Windows 控制台编码问题
if sys.platform == "win32":
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except (AttributeError, OSError):
        pass

import json
import os
import shutil
import sys
from copy import deepcopy
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Pt, Inches, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("错误：需要安装 python-docx。运行：python scripts/install_deps.py")
    sys.exit(1)


# ══════════════════════════════════════════════════════════════════
# 工具函数
# ══════════════════════════════════════════════════════════════════

def backup_file(filepath: str) -> str:
    """创建文件备份，返回备份路径."""
    backup_path = filepath + ".bak"
    # 如果备份已存在，追加时间戳
    if os.path.exists(backup_path):
        import time
        ts = time.strftime("%Y%m%d_%H%M%S")
        backup_path = f"{filepath}.bak.{ts}"
    shutil.copy2(filepath, backup_path)
    return backup_path


def cm_to_emu(cm: float) -> int:
    """厘米转 EMU."""
    return int(cm * 360000)


def pt_to_emu(pt: float) -> int:
    """磅转 EMU."""
    return int(pt * 12700)


def parse_alignment(align_str: str):
    """解析对齐方式字符串."""
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "居中": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
        "分散对齐": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    return mapping.get(align_str.lower(), None)


def set_run_font(run, font_config: dict):
    """设置 run 级别的字体属性，正确处理中文字体."""
    if "name_cn" in font_config and font_config["name_cn"]:
        run.font.name = font_config["name_cn"]
        # 设置东亚字体（中文）
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_config["name_cn"])
        if "name_en" in font_config and font_config["name_en"]:
            rFonts.set(qn('w:ascii'), font_config["name_en"])
            rFonts.set(qn('w:hAnsi'), font_config["name_en"])
    elif "name_en" in font_config and font_config["name_en"]:
        run.font.name = font_config["name_en"]

    if "size_pt" in font_config and font_config["size_pt"]:
        run.font.size = Pt(font_config["size_pt"])

    if "bold" in font_config:
        run.bold = font_config["bold"]
        # 同时设置 w:bCs（CJK/复杂脚本加粗），否则中文字不会显示加粗
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            run._element.insert(0, rPr)
        bCs = rPr.find(qn('w:bCs'))
        if font_config["bold"]:
            if bCs is None:
                bCs = parse_xml(f'<w:bCs {nsdecls("w")}/>')
                rPr.append(bCs)
        else:
            if bCs is not None:
                rPr.remove(bCs)
        # 同样设置 w:b 的显式值，确保不会被样式覆盖
        b = rPr.find(qn('w:b'))
        if font_config["bold"]:
            if b is None:
                b = parse_xml(f'<w:b {nsdecls("w")}/>')
                rPr.append(b)
        else:
            if b is not None:
                rPr.remove(b)

    if "italic" in font_config:
        run.italic = font_config["italic"]

    if "color" in font_config and font_config["color"]:
        hex_color = font_config["color"].lstrip("#")
        run.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )


# ══════════════════════════════════════════════════════════════════
# 格式修改函数
# ══════════════════════════════════════════════════════════════════

def apply_page_layout(doc: Document, page_config: dict):
    """应用页面布局设置."""
    for section in doc.sections:
        # 纸张大小
        paper_size = page_config.get("paper_size", "").upper()
        paper_sizes = {
            "A4": (21.0, 29.7),
            "LETTER": (21.59, 27.94),
            "LEGAL": (21.59, 35.56),
            "B5": (18.4, 26.0),
            "A5": (14.8, 21.0),
        }
        if paper_size in paper_sizes:
            w, h = paper_sizes[paper_size]
            section.page_width = Cm(w)
            section.page_height = Cm(h)

        # 页边距
        for key, section_attr in [
            ("margin_top_cm", "top_margin"),
            ("margin_bottom_cm", "bottom_margin"),
            ("margin_left_cm", "left_margin"),
            ("margin_right_cm", "right_margin"),
        ]:
            if key in page_config and page_config[key] is not None:
                setattr(section, section_attr, Cm(page_config[key]))


def apply_body_font(doc: Document, font_config: dict):
    """将字体设置应用到所有正文段落（非标题）."""
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        # 跳过标题段落
        if any(kw in style_name.lower() for kw in ["heading", "标题", "toc"]):
            continue
        for run in para.runs:
            set_run_font(run, font_config)


def apply_paragraph_format(doc: Document, para_config: dict):
    """应用段落格式设置."""
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if any(kw in style_name.lower() for kw in ["heading", "标题", "toc"]):
            continue

        pf = para.paragraph_format

        # 行间距
        if "line_spacing" in para_config and para_config["line_spacing"] is not None:
            ls = para_config["line_spacing"]
            # 支持 "1.5", "2.0" 等倍数
            pf.line_spacing = float(ls)

        # 段前段后间距
        if "space_before_pt" in para_config:
            pf.space_before = Pt(para_config["space_before_pt"])
        if "space_after_pt" in para_config:
            pf.space_after = Pt(para_config["space_after_pt"])

        # 对齐方式
        if "alignment" in para_config and para_config["alignment"]:
            align = parse_alignment(para_config["alignment"])
            if align is not None:
                para.alignment = align

        # 首行缩进（字符数）
        if "first_line_indent_chars" in para_config:
            chars = para_config["first_line_indent_chars"]
            # 首行缩进按字符数计算：假设当前字号对应的字符宽度
            # 默认按 12pt 字号的 2 字符 ≈ 0.74cm
            indent_cm = chars * 0.37  # 近似值
            pf.first_line_indent = Cm(indent_cm)

        # 首行缩进（精确厘米值，优先级高于字符数）
        if "first_line_indent_cm" in para_config:
            pf.first_line_indent = Cm(para_config["first_line_indent_cm"])


def apply_heading_styles(doc: Document, headings_config: dict):
    """应用标题样式."""
    heading_mapping = {
        "h1": "Heading 1",
        "h2": "Heading 2",
        "h3": "Heading 3",
        "h4": "Heading 4",
        "h5": "Heading 5",
        "h6": "Heading 6",
    }

    for h_key, font_config in headings_config.items():
        heading_name = heading_mapping.get(h_key, h_key)

        # 修改 Word 内置标题样式
        try:
            style = doc.styles[heading_name]
        except KeyError:
            print(f"  ⚠️  样式 '{heading_name}' 不存在，跳过")
            continue

        if "name_cn" in font_config or "name_en" in font_config:
            style.font.name = font_config.get("name_cn") or font_config.get("name_en")
            # 设置东亚字体
            rPr = style.element.find(qn('w:rPr'))
            if rPr is None:
                rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                style.element.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
                rPr.insert(0, rFonts)
            if "name_cn" in font_config and font_config["name_cn"]:
                rFonts.set(qn('w:eastAsia'), font_config["name_cn"])
            if "name_en" in font_config and font_config["name_en"]:
                rFonts.set(qn('w:ascii'), font_config["name_en"])
                rFonts.set(qn('w:hAnsi'), font_config["name_en"])

        if "size_pt" in font_config:
            style.font.size = Pt(font_config["size_pt"])
        if "bold" in font_config:
            style.font.bold = font_config["bold"]
            # 同样在样式 XML 中设置 w:bCs，确保中文字加粗生效
            s_rPr = style.element.find(qn('w:rPr'))
            if s_rPr is None:
                s_rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                style.element.append(s_rPr)
            sbCs = s_rPr.find(qn('w:bCs'))
            if font_config["bold"]:
                if sbCs is None:
                    sbCs = parse_xml(f'<w:bCs {nsdecls("w")}/>')
                    s_rPr.append(sbCs)
            else:
                if sbCs is not None:
                    s_rPr.remove(sbCs)
            # 同样显式设置 w:b
            sb = s_rPr.find(qn('w:b'))
            if font_config["bold"]:
                if sb is None:
                    sb = parse_xml(f'<w:b {nsdecls("w")}/>')
                    s_rPr.append(sb)
            else:
                if sb is not None:
                    s_rPr.remove(sb)
        if "color" in font_config and font_config["color"]:
            hex_color = font_config["color"].lstrip("#")
            style.font.color.rgb = RGBColor(
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16),
            )

    # 同时也直接修改已有标题段落中的 run（覆盖样式未生效的情况）
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        for h_key, heading_name in heading_mapping.items():
            if heading_name.lower() == style_name.lower() and h_key in headings_config:
                for run in para.runs:
                    set_run_font(run, headings_config[h_key])


def apply_header_footer(doc: Document, hf_config: dict):
    """应用页眉页脚设置."""
    for section in doc.sections:
        # 页眉
        if "header_text" in hf_config and hf_config["header_text"]:
            header = section.header
            header.is_linked_to_previous = False
            if header.paragraphs:
                header.paragraphs[0].text = hf_config["header_text"]
            else:
                para = header.add_paragraph()
                para.text = hf_config["header_text"]

        # 页脚
        if "footer_text" in hf_config and hf_config["footer_text"]:
            footer = section.footer
            footer.is_linked_to_previous = False
            if footer.paragraphs:
                footer.paragraphs[0].text = hf_config["footer_text"]
            else:
                para = footer.add_paragraph()
                para.text = hf_config["footer_text"]

        # 页码
        if hf_config.get("show_page_number", False):
            footer = section.footer
            footer.is_linked_to_previous = False
            position = hf_config.get("page_number_position", "center")
            # 添加页码域代码
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            if position == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif position == "right":
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # 添加 PAGE 域
            run = para.add_run()
            fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._element.append(fldChar_begin)

            run2 = para.add_run()
            instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run2._element.append(instrText)

            run3 = para.add_run()
            fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run3._element.append(fldChar_end)


def remove_empty_paragraphs(doc: Document):
    """清理连续的空段落."""
    removed = 0
    for para in list(doc.paragraphs)[1:]:  # 保留第一个段落
        if not para.text.strip() and not para.runs:
            p = para._element
            p.getparent().remove(p)
            removed += 1
    return removed


# ══════════════════════════════════════════════════════════════════
# 模版模式
# ══════════════════════════════════════════════════════════════════

def extract_template_styles(template_path: str) -> dict:
    """从模版 .docx 提取格式配置."""
    doc = Document(template_path)
    config = {}

    # 提取页面设置
    if doc.sections:
        section = doc.sections[0]
        width_cm = section.page_width / 360000
        height_cm = section.page_height / 360000

        # 识别纸张大小
        paper_sizes = {
            (21.0, 29.7): "A4",
            (21.59, 27.94): "Letter",
            (21.59, 35.56): "Legal",
        }
        for (w, h), name in paper_sizes.items():
            if abs(width_cm - w) < 0.3 and abs(height_cm - h) < 0.3:
                config.setdefault("page", {})["paper_size"] = name
                break

        config.setdefault("page", {}).update({
            "margin_top_cm": round(section.top_margin / 360000, 2),
            "margin_bottom_cm": round(section.bottom_margin / 360000, 2),
            "margin_left_cm": round(section.left_margin / 360000, 2),
            "margin_right_cm": round(section.right_margin / 360000, 2),
        })

    # 提取标题样式
    heading_mapping = {
        "Heading 1": "h1", "Heading 2": "h2", "Heading 3": "h3",
        "Heading 4": "h4", "Heading 5": "h5", "Heading 6": "h6",
        "标题 1": "h1", "标题 2": "h2", "标题 3": "h3",
    }

    headings_config = {}
    for style_name, h_key in heading_mapping.items():
        try:
            style = doc.styles[style_name]
            if style.font.name or style.font.size:
                headings_config[h_key] = {}
                if style.font.name:
                    headings_config[h_key]["name_cn"] = style.font.name
                    headings_config[h_key]["name_en"] = style.font.name
                if style.font.size:
                    headings_config[h_key]["size_pt"] = round(style.font.size / 12700, 1)
                if style.font.bold is not None:
                    headings_config[h_key]["bold"] = style.font.bold
        except KeyError:
            pass

    if headings_config:
        config["headings"] = headings_config

    # 提取正文字体（采样前几个段落）
    font_names = []
    font_sizes = []
    for para in doc.paragraphs[:20]:
        style_name = para.style.name if para.style else ""
        if "heading" in style_name.lower() or "标题" in style_name:
            continue
        for run in para.runs:
            if run.font.name:
                font_names.append(run.font.name)
            if run.font.size:
                font_sizes.append(run.font.size / 12700)

    if font_names or font_sizes:
        config["body_font"] = {}
        if font_names:
            config["body_font"]["name_cn"] = max(set(font_names), key=font_names.count)
            config["body_font"]["name_en"] = max(set(font_names), key=font_names.count)
        if font_sizes:
            config["body_font"]["size_pt"] = round(sum(font_sizes) / len(font_sizes), 1)

    return config


# ══════════════════════════════════════════════════════════════════
# 预设格式
# ══════════════════════════════════════════════════════════════════

PRESETS = {
    "gb-thesis": {
        "description": "GB/T 学位论文标准格式",
        "page": {
            "paper_size": "A4",
            "margin_top_cm": 3.0,
            "margin_bottom_cm": 3.0,
            "margin_left_cm": 3.0,
            "margin_right_cm": 2.5,
        },
        "body_font": {
            "name_cn": "宋体",
            "name_en": "Times New Roman",
            "size_pt": 12,
        },
        "paragraph": {
            "line_spacing": 1.5,
            "first_line_indent_chars": 2,
            "alignment": "justify",
        },
        "headings": {
            "h1": {"name_cn": "黑体", "size_pt": 16, "bold": True},
            "h2": {"name_cn": "黑体", "size_pt": 14, "bold": True},
            "h3": {"name_cn": "黑体", "size_pt": 12, "bold": True},
        },
        "header_footer": {
            "show_page_number": True,
            "page_number_position": "center",
        },
    },
    "ieee": {
        "description": "IEEE 会议/期刊论文格式",
        "page": {
            "paper_size": "Letter",
            "margin_top_cm": 1.78,
            "margin_bottom_cm": 1.78,
            "margin_left_cm": 1.78,
            "margin_right_cm": 1.78,
        },
        "body_font": {
            "name_en": "Times New Roman",
            "size_pt": 10,
        },
        "paragraph": {
            "line_spacing": 1.0,
            "alignment": "justify",
        },
    },
    "apa": {
        "description": "APA 7th 论文格式",
        "page": {
            "paper_size": "Letter",
            "margin_top_cm": 2.54,
            "margin_bottom_cm": 2.54,
            "margin_left_cm": 2.54,
            "margin_right_cm": 2.54,
        },
        "body_font": {
            "name_en": "Times New Roman",
            "size_pt": 12,
        },
        "paragraph": {
            "line_spacing": 2.0,
            "first_line_indent_chars": 5,  # 0.5 inch for APA
            "alignment": "left",
        },
    },
}


# ══════════════════════════════════════════════════════════════════
# 主程序
# ══════════════════════════════════════════════════════════════════

def load_config(config_path: str = None, template_path: str = None,
                preset: str = None) -> dict:
    """加载格式配置."""
    if preset and preset in PRESETS:
        print(f"📋 使用预设格式：{PRESETS[preset]['description']}")
        return PRESETS[preset]

    if template_path:
        print(f"📋 从模版提取格式：{template_path}")
        config = extract_template_styles(template_path)
        print(f"   提取到的配置：")
        print(json.dumps(config, ensure_ascii=False, indent=2))
        return config

    if config_path:
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)

    print("错误：必须指定 --config、--template 或 --preset 之一")
    sys.exit(1)


def apply_format(input_path: str, output_path: str, config: dict) -> list:
    """应用格式配置到文档，返回修改摘要列表."""
    doc = Document(input_path)
    changes = []

    # 1. 页面布局
    if "page" in config:
        apply_page_layout(doc, config["page"])
        changes.append(f"✅ 页面布局：{json.dumps(config['page'], ensure_ascii=False)}")

    # 2. 标题样式
    if "headings" in config:
        apply_heading_styles(doc, config["headings"])
        for hk, hv in config["headings"].items():
            changes.append(f"✅ 标题样式 {hk}：字体={hv.get('name_cn', hv.get('name_en', '-'))} "
                           f"字号={hv.get('size_pt', '-')}pt 加粗={hv.get('bold', False)}")

    # 3. 正文字体
    if "body_font" in config:
        apply_body_font(doc, config["body_font"])
        changes.append(f"✅ 正文字体：{json.dumps(config['body_font'], ensure_ascii=False)}")

    # 4. 段落格式
    if "paragraph" in config:
        apply_paragraph_format(doc, config["paragraph"])
        changes.append(f"✅ 段落格式：{json.dumps(config['paragraph'], ensure_ascii=False)}")

    # 5. 页眉页脚
    if "header_footer" in config:
        apply_header_footer(doc, config["header_footer"])
        changes.append(f"✅ 页眉页脚已更新")

    # 6. 保存
    doc.save(output_path)
    changes.append(f"\n📁 输出文件：{output_path}")
    return changes


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="论文格式修改引擎 — Paper Formatter",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例：
  # 预设格式
  python format_docx.py paper.docx output.docx --preset gb-thesis

  # JSON 配置文件
  python format_docx.py paper.docx output.docx --config my_format.json

  # 从模版学习
  python format_docx.py paper.docx output.docx --template template.docx

  # 列出预设
  python format_docx.py --list-presets
        """
    )
    parser.add_argument("input", nargs="?", help="输入 .docx 文件路径")
    parser.add_argument("output", nargs="?", help="输出 .docx 文件路径")
    parser.add_argument("--config", "-c", help="JSON 格式配置文件路径")
    parser.add_argument("--template", "-t", help="模版 .docx 文件路径（提取样式并应用）")
    parser.add_argument("--preset", "-p", choices=list(PRESETS.keys()),
                        help="使用预设格式")
    parser.add_argument("--list-presets", action="store_true",
                        help="列出所有预设格式")
    parser.add_argument("--no-backup", action="store_true",
                        help="不创建备份（谨慎使用）")

    args = parser.parse_args()

    # 列出预设
    if args.list_presets:
        print("📋 可用预设格式：\n")
        for name, preset in PRESETS.items():
            print(f"  {name}: {preset['description']}")
        return

    # 验证参数
    if not args.input:
        parser.print_help()
        print("\n错误：请指定输入文件")
        sys.exit(1)

    if not os.path.exists(args.input):
        print(f"错误：文件不存在 - {args.input}")
        sys.exit(1)

    if not args.output:
        # 默认输出文件名
        base, ext = os.path.splitext(args.input)
        args.output = f"{base}_formatted{ext}"

    if not any([args.config, args.template, args.preset]):
        print("错误：必须指定 --config、--template 或 --preset 之一")
        print("💡 提示：运行 --list-presets 查看可用预设")
        sys.exit(1)

    # --- 执行 ---
    print("=" * 60)
    print("  📝 Paper Formatter — 论文格式修改引擎")
    print("=" * 60)

    # 1. 备份
    if not args.no_backup:
        backup_path = backup_file(args.input)
        print(f"\n📦 备份已创建：{backup_path}")

    # 2. 加载配置
    config = load_config(
        config_path=args.config,
        template_path=args.template,
        preset=args.preset,
    )

    # 3. 应用格式
    print(f"\n🔧 正在修改格式...")
    changes = apply_format(args.input, args.output, config)

    # 4. 输出摘要
    print(f"\n📊 修改摘要：")
    for change in changes:
        print(f"   {change}")

    print(f"\n✅ 格式修改完成！")
    if not args.no_backup:
        print(f"💾 如需恢复原文件，备份位于：{backup_path}")
    print()


if __name__ == "__main__":
    main()
